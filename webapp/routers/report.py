from fastapi import APIRouter
from fastapi.responses import FileResponse
from pydantic import BaseModel
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from datetime import datetime
import os
import uuid
from collections import Counter
from PIL import Image as PILImage
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from matplotlib.patches import Patch

router = APIRouter(prefix="/report", tags=["Report"])

class Vulnerability(BaseModel):
    id: int
    title: str
    severity: str
    cvss_score: str
    cvss_vector: str
    description: str
    recommendation: str
    reference: str
    instanceId: str

class ReportRequest(BaseModel):
    app_title: str
    scope: str
    urls: str
    analyst_name: str
    requester_name: str
    vulnerabilities: list[Vulnerability]
    evidence_data: dict

def generate_pie_chart(vulns, chart_path):
    severity_order = ["Critical", "High", "Medium", "Low"]
    color_map = {
        "Critical": "maroon",
        "High": "red",
        "Medium": "orange",
        "Low": "green"
    }

    severity_counts = Counter(v.severity for v in vulns)
    labels = [sev for sev in severity_order if severity_counts[sev] > 0]
    sizes = [severity_counts[sev] for sev in labels]
    colors = [color_map[sev] for sev in labels]

    def make_autopct(values):
        def _autopct(pct):
            total = sum(values)
            val = int(round(pct * total / 100.0))
            return f"{val}" if val > 0 else ""
        return _autopct

    plt.figure(figsize=(4, 4))
    wedges, texts, autotexts = plt.pie(
        sizes,
        labels=labels,
        colors=colors,
        autopct=make_autopct(sizes),
        startangle=140
    )
    for autotext in autotexts:
        autotext.set_color('white')
        autotext.set_fontsize(10)
        autotext.set_weight("bold")

    legend_labels = [f"{severity_counts[sev]} {sev}" for sev in labels]
    patches = [Patch(color=color_map[sev], label=legend_labels[i]) for i, sev in enumerate(labels)]
    plt.legend(handles=patches, loc="best")
    plt.axis('equal')
    plt.tight_layout()
    plt.savefig(chart_path)
    plt.close()

@router.post("/", response_class=FileResponse)
def generate_report(payload: ReportRequest):
    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)

    severity_order = ["Critical", "High", "Medium", "Low"]
    severity_colors = {
        "Critical": RGBColor(128, 0, 0),
        "High": RGBColor(255, 0, 0),
        "Medium": RGBColor(255, 191, 0),
        "Low": RGBColor(0, 128, 0)
    }

    logo_path = "uploaded_logos/logo.png"
    if os.path.exists(logo_path):
        para = doc.add_paragraph()
        para.alignment = 1
        para.add_run().add_picture(logo_path, width=Inches(2.5))

    title = doc.add_paragraph("\n\nDynamic Vulnerability Assessment")
    title.alignment = 1
    run = title.runs[0]
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0, 102, 204)

    app = doc.add_paragraph(payload.app_title)
    app.alignment = 1
    run = app.runs[0]
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0, 102, 204)

    requester = doc.add_paragraph()
    requester.alignment = 0
    req_run = requester.add_run("Requested by: ")
    req_run.bold = True
    requester.add_run(payload.requester_name)

    doc.add_page_break()

    section = doc.sections[0]
    if os.path.exists(logo_path):
        header = section.header
        header_para = header.paragraphs[0]
        header_para.alignment = 2
        header_para.add_run().add_picture(logo_path, width=Inches(1.0))

    footer = section.footer.paragraphs[0]
    footer.text = "Confidential - For Internal Use Only"
    footer.alignment = 1

    doc.add_heading("Version Information", level=2).runs[0].font.size = Pt(18)
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    headers = ["Date", "Application Version", "Reviewer"]
    for i, h in enumerate(headers):
        table.cell(0, i).text = h
    today = datetime.now().strftime("%d-%b-%Y")
    table.cell(1, 0).text = today
    table.cell(1, 1).text = "Initial Draft"
    table.cell(1, 2).text = payload.analyst_name
    table.cell(2, 0).text = today
    table.cell(2, 1).text = "Peer Review"
    table.cell(3, 0).text = today
    table.cell(3, 1).text = "Approved"

    doc.add_page_break()

    os.makedirs("generated_reports", exist_ok=True)
    pie_path = os.path.join("generated_reports", f"pie_{uuid.uuid4().hex}.png")
    generate_pie_chart(payload.vulnerabilities, pie_path)
    doc.add_paragraph().add_run("Vulnerability Severity Distribution").bold = True
    doc.add_picture(pie_path, width=Inches(4.5))
    doc.paragraphs[-1].alignment = 1
    doc.add_page_break()

    doc.add_heading("Summary Table", level=2).runs[0].font.size = Pt(18)
    summary_table = doc.add_table(rows=1, cols=4)
    summary_table.style = 'Table Grid'
    hdrs = ["Sl. No.", "Security Observation", "Risk Rating", "Page No."]
    for i, h in enumerate(hdrs):
        summary_table.cell(0, i).text = h

    count = 1
    page_counter = 4
    for sev in severity_order:
        group = [v for v in payload.vulnerabilities if v.severity == sev]
        if group:
            row = summary_table.add_row().cells
            row[0].merge(row[3])
            heading = row[0].paragraphs[0].add_run(f"{sev} Severity")
            heading.bold = True
            heading.font.color.rgb = severity_colors[sev]
            for v in group:
                row = summary_table.add_row().cells
                row[0].text = str(count)
                row[1].text = v.title
                risk = row[2].paragraphs[0].add_run(v.severity)
                risk.font.color.rgb = severity_colors[v.severity]
                row[3].text = str(page_counter)
                count += 1
                page_counter += 1

    doc.add_page_break()
    doc.add_heading("URLs and Scope", level=2).runs[0].font.size = Pt(18)
    doc.add_paragraph(f"URLs: {payload.urls}")
    doc.add_paragraph(f"Scope: {payload.scope}")
    doc.add_page_break()

    doc.add_heading("Vulnerability Details", level=2).runs[0].font.size = Pt(18)

    for idx, vuln in enumerate(payload.vulnerabilities, 1):
        doc.add_page_break()
        doc.add_heading(f"{idx}. {vuln.title}", level=3).runs[0].font.size = Pt(18)

        para = doc.add_paragraph()
        para.add_run("Severity: ").bold = True
        sev_run = para.add_run(vuln.severity)
        sev_run.bold = True
        sev_run.font.color.rgb = severity_colors.get(vuln.severity, RGBColor(0, 0, 0))

        doc.add_paragraph(f"CVSS Score: {vuln.cvss_score}")
        doc.add_paragraph(f"CVSS Vector: {vuln.cvss_vector}")

        doc.add_heading("Description", level=4).runs[0].font.size = Pt(18)
        doc.add_paragraph(vuln.description)

        doc.add_heading("Evidence", level=4).runs[0].font.size = Pt(18)
        steps = payload.evidence_data.get(str(vuln.instanceId or vuln.id), [])
        for step_idx, step in enumerate(steps, 1):
            doc.add_paragraph(f"Step {step_idx}: {step.get('comment', '')}")

            paths = step.get("screenshotPath", [])
            if isinstance(paths, str):
                paths = [paths]

            for img_path in paths:
                clean_path = img_path.lstrip("/")
                full_path = os.path.join(".", clean_path)
                if os.path.exists(full_path):
                    try:
                        with PILImage.open(full_path) as img:
                            img.verify()
                        doc.add_picture(full_path, width=Inches(4))
                    except Exception as e:
                        doc.add_paragraph(f"[Invalid image: {clean_path}]")
                else:
                    doc.add_paragraph(f"[Image not found: {clean_path}]")

        doc.add_heading("Recommendation", level=4).runs[0].font.size = Pt(18)
        doc.add_paragraph(vuln.recommendation)
        doc.add_heading("Reference", level=4).runs[0].font.size = Pt(18)
        doc.add_paragraph(vuln.reference)

    filename = f"report_{uuid.uuid4().hex}.docx"
    filepath = os.path.join("generated_reports", filename)
    doc.save(filepath)

    return FileResponse(
        filepath,
        filename="DVA_Report.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
