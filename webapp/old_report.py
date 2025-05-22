from fastapi import APIRouter, HTTPException
from fastapi.responses import FileResponse
from pydantic import BaseModel
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from datetime import datetime
import matplotlib.pyplot as plt
from collections import Counter
from matplotlib.patches import Patch
import os
import uuid
from PIL import Image as PILImage
from docx2pdf import convert

router = APIRouter(prefix="/report", tags=["Report"])

class Vulnerability(BaseModel):
    id: int
    title: str
    severity: str
    cvss_score: str
    cvss_vector: str
    description: str
    recommendation: str
    reference: str
    instanceId: str

class ReportRequest(BaseModel):
    app_title: str
    scope: str
    urls: str
    analyst_name: str
    requester_name: str
    vulnerabilities: list[Vulnerability]
    evidence_data: dict

def generate_pie_chart(vulns, chart_path):
    print("✅ Using UPDATED pie chart logic")

    severity_order = ["Critical", "High", "Medium", "Low"]
    color_map = {
        "Critical": "maroon",
        "High": "red",
        "Medium": "orange",
        "Low": "green"
    }

    severity_counts = Counter(v.severity for v in vulns)
    labels = [sev for sev in severity_order if severity_counts[sev] > 0]
    sizes = [severity_counts[sev] for sev in labels]
    colors = [color_map[sev] for sev in labels]

    def make_autopct(values):
        def _autopct(pct):
            total = sum(values)
            val = int(round(pct * total / 100.0))
            return f"{val}" if val > 0 else ""
        return _autopct

    plt.figure(figsize=(4, 4))
    wedges, texts, autotexts = plt.pie(
        sizes,
        labels=labels,
        colors=colors,
        autopct=make_autopct(sizes),
        startangle=140
    )

    for autotext in autotexts:
        autotext.set_color('white')
        autotext.set_fontsize(10)
        autotext.set_weight("bold")

    legend_labels = [f"{severity_counts[sev]} {sev}" for sev in labels]
    patches = [Patch(color=color_map[sev], label=legend_labels[i]) for i, sev in enumerate(labels)]
    plt.legend(handles=patches, loc="best")
    plt.axis('equal')
    plt.tight_layout()
    plt.savefig(chart_path)
    plt.close()

def build_report_doc(payload: ReportRequest, output_path: str):
    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)

    severity_order = ["Critical", "High", "Medium", "Low"]
    severity_colors = {
        "Critical": RGBColor(128, 0, 0),
        "High": RGBColor(255, 0, 0),
        "Medium": RGBColor(255, 191, 0),
        "Low": RGBColor(0, 128, 0)
    }

    logo_path = "uploaded_logos/logo.png"
    if os.path.exists(logo_path):
        doc.add_paragraph().add_run().add_picture(logo_path, width=Inches(2.5)).paragraph.alignment = 1

    title = doc.add_paragraph("\n\nDynamic Vulnerability Assessment")
    title.alignment = 1
    run = title.runs[0]
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0, 102, 204)

    app = doc.add_paragraph(payload.app_title)
    app.alignment = 1
    run = app.runs[0]
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0, 102, 204)

    doc.add_paragraph(f"Requested by: {payload.requester_name}")
    doc.add_page_break()

    section = doc.sections[0]
    if os.path.exists(logo_path):
        section.header.paragraphs[0].add_run().add_picture(logo_path, width=Inches(1.0)).paragraph.alignment = 2
    section.footer.paragraphs[0].text = "Confidential - For Internal Use Only"
    section.footer.paragraphs[0].alignment = 1

    doc.add_heading("Version Information", level=2)
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    headers = ["Date", "Application Version", "Reviewer"]
    for i, h in enumerate(headers):
        table.cell(0, i).text = h
    today = datetime.now().strftime("%d-%b-%Y")
    table.cell(1, 0).text = today
    table.cell(1, 1).text = "Initial Draft"
    table.cell(1, 2).text = payload.analyst_name
    table.cell(2, 0).text = today
    table.cell(2, 1).text = "Peer Review"
    table.cell(3, 0).text = today
    table.cell(3, 1).text = "Approved"

    chart_path = output_path.replace(".docx", "_chart.png")
    try:
        generate_pie_chart(payload.vulnerabilities, chart_path)
        doc.add_paragraph("\n")
        doc.add_picture(chart_path, width=Inches(4.5))
    except Exception as e:
        doc.add_paragraph(f"[Failed to generate chart: {str(e)}]")

    doc.add_page_break()
    doc.add_heading("Summary Table", level=2)
    summary_table = doc.add_table(rows=1, cols=4)
    summary_table.style = 'Table Grid'
    hdrs = ["Sl. No.", "Security Observation", "Risk Rating", "Page No."]
    for i, h in enumerate(hdrs):
        summary_table.cell(0, i).text = h

    count = 1
    page_counter = 4
    for sev in severity_order:
        group = [v for v in payload.vulnerabilities if v.severity == sev]
        if group:
            row = summary_table.add_row().cells
            row[0].merge(row[3])
            heading = row[0].paragraphs[0].add_run(f"{sev} Severity")
            heading.bold = True
            heading.font.color.rgb = severity_colors[sev]
            for v in group:
                row = summary_table.add_row().cells
                row[0].text = str(count)
                row[1].text = v.title
                risk = row[2].paragraphs[0].add_run(v.severity)
                risk.font.color.rgb = severity_colors[v.severity]
                row[3].text = str(page_counter)
                count += 1
                page_counter += 1

    doc.add_page_break()
    doc.add_heading("URLs and Scope", level=2)
    doc.add_paragraph(f"URLs: {payload.urls}")
    doc.add_paragraph(f"Scope: {payload.scope}")
    doc.add_page_break()

    doc.add_heading("Vulnerability Details", level=2)
    for idx, vuln in enumerate(payload.vulnerabilities, 1):
        doc.add_page_break()
        doc.add_heading(f"{idx}. {vuln.title}", level=3)

        p = doc.add_paragraph()
        p.add_run("Severity: ").bold = True
        sev_run = p.add_run(vuln.severity)
        sev_run.bold = True
        sev_run.font.color.rgb = severity_colors.get(vuln.severity, RGBColor(0, 0, 0))

        doc.add_paragraph(f"CVSS Score: {vuln.cvss_score}")
        doc.add_paragraph(f"CVSS Vector: {vuln.cvss_vector}")
        doc.add_heading("Description", level=4)
        doc.add_paragraph(vuln.description)

        doc.add_heading("Evidence", level=4)
        steps = payload.evidence_data.get(str(vuln.instanceId or vuln.id), {}).get("steps", [])
        print(f"🔍 Evidence for {vuln.title}: {steps}")
        for step_idx, step in enumerate(steps, 1):
            doc.add_paragraph(f"Step {step_idx}", style="List Number")
            if step["type"] == "text":
                doc.add_paragraph(step["content"])
            elif step["type"] == "image":
                content = step["content"]
                print(f"🖼️ Step {step_idx} image content: {content}")
                if isinstance(content, str):
                    content = [content]
                for img_path in content:
                    filename = os.path.basename(img_path)
                    full_path = os.path.join("uploaded_evidence", filename)
                    if os.path.exists(full_path):
                        try:
                            with PILImage.open(full_path) as img:
                                img.verify()
                            doc.add_picture(full_path, width=Inches(4))
                        except Exception as e:
                            doc.add_paragraph(f"[Invalid or unreadable image: {filename}]")
                            print(f"❌ PIL failed to open {filename}: {e}")
                    else:
                        doc.add_paragraph(f"[Missing Image: {filename}]")
                        print(f"❌ Image not found: {full_path}")

        doc.add_heading("Recommendation", level=4)
        doc.add_paragraph(vuln.recommendation)
        doc.add_heading("Reference", level=4)
        doc.add_paragraph(vuln.reference)

    doc.save(output_path)

@router.post("/word")
def generate_word_report(payload: ReportRequest):
    try:
        os.makedirs("generated_reports", exist_ok=True)
        safe_name = payload.app_title.replace(" ", "_")
        date_str = datetime.now().strftime("%m%d%y")
        output_path = os.path.join("generated_reports", f"{safe_name}_ManualReport_{date_str}.docx")
        print(f"📝 Starting report generation: {output_path}")
        build_report_doc(payload, output_path)
        print(f"✅ Report successfully saved to {output_path}")
        return FileResponse(output_path, filename=os.path.basename(output_path),
                            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    except Exception as e:
        print(f"❌ Report generation failed: {e}")
        raise HTTPException(status_code=500, detail=f"Report generation failed: {str(e)}")


@router.post("/pdf")
def generate_pdf_report(payload: ReportRequest):
    os.makedirs("generated_reports", exist_ok=True)
    safe_name = payload.app_title.replace(" ", "_")
    date_str = datetime.now().strftime("%m%d%y")
    word_path = os.path.join("generated_reports", f"{safe_name}_ManualReport_{date_str}.docx")
    pdf_path = word_path.replace(".docx", ".pdf")
    build_report_doc(payload, word_path)
    try:
        convert(word_path, pdf_path)
        return FileResponse(pdf_path, filename=os.path.basename(pdf_path), media_type="application/pdf")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"PDF generation failed: {str(e)}")
