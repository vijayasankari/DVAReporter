from PyQt5.QtWidgets import QWidget, QVBoxLayout, QPushButton, QFileDialog, QMessageBox
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import os


def set_cell_background(cell, color="000000"):
    tc_pr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tc_pr.append(shd)


class GenerateTab(QWidget):
    def __init__(self, project_tab, summary_tab, logo_path):
        super().__init__()
        self.project_tab = project_tab
        self.summary_tab = summary_tab
        self.logo_path = logo_path

        layout = QVBoxLayout()
        self.generate_button = QPushButton("Generate Word Report")
        self.generate_button.clicked.connect(self.generate_report)
        layout.addWidget(self.generate_button)
        self.setLayout(layout)

    def generate_report(self):
        file_path, _ = QFileDialog.getSaveFileName(self, "Save Report", "", "Word Documents (*.docx)")
        if not file_path:
            return

        doc = Document()
        doc.styles['Normal'].font.name = 'Calibri'
        doc.styles['Normal'].font.size = Pt(11)

        # === First Page ===
        if self.logo_path and os.path.exists(self.logo_path):
            para = doc.add_paragraph()
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            para.add_run().add_picture(self.logo_path, width=Inches(2.5))

        title = doc.add_paragraph()
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = title.add_run("\n\nDynamic Vulnerability Assessment\n")
        run.bold = True
        run.font.size = Pt(28)
        run.font.color.rgb = RGBColor(0, 102, 204)

        app = doc.add_paragraph()
        app.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = app.add_run(self.project_tab.get_data().get("project_title", "Application"))
        run.bold = True
        run.font.size = Pt(24)
        run.font.color.rgb = RGBColor(0, 102, 204)

        doc.add_page_break()

        # === Header Logo ===
        section = doc.sections[0]
        header = section.header
        if self.logo_path and os.path.exists(self.logo_path):
            para = header.paragraphs[0]
            para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            para.add_run().add_picture(self.logo_path, width=Inches(1.5))

        # === Version Information ===
        self.add_heading(doc, "Version Information")
        table = doc.add_table(rows=4, cols=3)
        table.style = 'Table Grid'
        headers = ["Date", "Application Version", "Reviewer"]
        for i, text in enumerate(headers):
            cell = table.cell(0, i)
            run = cell.paragraphs[0].add_run(text)
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            set_cell_background(cell)

        today = datetime.now().strftime("%d-%b-%Y")
        analyst = self.project_tab.get_data().get("analyst_name", "")
        table.cell(1, 0).text = today
        table.cell(1, 1).text = "Initial Draft"
        table.cell(1, 2).text = analyst
        table.cell(2, 0).text = today
        table.cell(2, 1).text = "Peer Review"
        table.cell(2, 2).text = ""
        table.cell(3, 0).text = today
        table.cell(3, 1).text = "Approved"
        table.cell(3, 2).text = ""

        doc.add_page_break()

        # === Summary Table ===
        self.add_heading(doc, "Summary Table")
        summary = self.summary_tab.get_summary_data()
        severity_order = ["Critical", "High", "Medium", "Low"]
        severity_colors = {
            "Critical": RGBColor(128, 0, 0),
            "High": RGBColor(255, 0, 0),
            "Medium": RGBColor(255, 191, 0),
            "Low": RGBColor(0, 128, 0)
        }

        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        headers = ["Sl. No.", "Security Observation", "Risk Rating", "Page No."]
        for i, h in enumerate(headers):
            cell = table.cell(0, i)
            run = cell.paragraphs[0].add_run(h)
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            set_cell_background(cell)

        counter = 1
        for sev in severity_order:
            group = [s for s in summary if s["severity"] == sev]
            if group:
                row = table.add_row().cells
                row[0].merge(row[3])
                run = row[0].paragraphs[0].add_run(f"{sev} Severity")
                run.bold = True
                run.font.color.rgb = severity_colors[sev]

                for item in group:
                    row = table.add_row().cells
                    row[0].text = str(counter)
                    row[1].text = item["title"]
                    run = row[2].paragraphs[0].add_run(item["severity"])
                    run.font.color.rgb = severity_colors[item["severity"]]
                    row[3].text = str(counter + 3)
                    counter += 1

        # === URLs and Scope ===
        doc.add_page_break()
        self.add_heading(doc, "URLs and Scope")
        doc.add_paragraph("URLs: https://example.com")
        scope_text = self.project_tab.get_data().get("scope", "")
        doc.add_paragraph(f"Scope: {scope_text}")

        # === Vulnerability Details ===
        doc.add_page_break()
        self.add_heading(doc, "Vulnerability Details")

        grouped_vulns = []
        for sev in severity_order:
            grouped_vulns.extend([v for v in self.summary_tab.vulns if v["severity"] == sev])

        for i, vuln in enumerate(grouped_vulns, 1):
            doc.add_page_break()
            heading_para = doc.add_paragraph()
            heading_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            title_run = heading_para.add_run(f"{i}. {vuln['title']}")
            title_run.bold = True
            title_run.font.size = Pt(18)

            p = doc.add_paragraph()
            p.add_run("Severity: ").bold = True
            r = p.add_run(vuln["severity"])
            r.bold = True
            r.font.color.rgb = severity_colors.get(vuln["severity"], RGBColor(0, 102, 204))

            p = doc.add_paragraph()
            p.add_run("CVSS Score: ").bold = True
            p.add_run(vuln.get("cvss_score", ""))

            p = doc.add_paragraph()
            p.add_run("CVSS Vector: ").bold = True
            p.add_run(vuln.get("cvss_vector", ""))

            self.add_heading(doc, "Description", level=4)
            doc.add_paragraph(vuln.get("description", ""))

            self.add_heading(doc, "Evidence", level=4)
            doc.add_paragraph(vuln.get("evidence", ""))

            self.add_heading(doc, "Recommendation", level=4)
            doc.add_paragraph(vuln.get("recommendation", ""))

            self.add_heading(doc, "Reference", level=4)
            doc.add_paragraph(vuln.get("reference", ""))

        doc.save(file_path)
        QMessageBox.information(self, "Success", f"Report saved to:\n{file_path}")

    def add_heading(self, doc, text, level=2):
        heading = doc.add_heading(level=level)
        run = heading.add_run(text)
        run.font.size = Pt(16)
